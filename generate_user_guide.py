"""
Run once to generate 'How to Use ProfTracker.docx' in the project root.
Usage:  python generate_user_guide.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD   = RGBColor(0xF4, 0xD6, 0x41)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x88, 0x88, 0x99)
DARK   = RGBColor(0x0C, 0x0C, 0x14)
LIGHT  = RGBColor(0xDC, 0xDC, 0xE8)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = GOLD if level == 1 else LIGHT
        run.font.bold = True
        run.font.name = "Segoe UI"
        run.font.size = Pt(18 if level == 1 else 13)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text, bold=False, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name  = "Segoe UI"
    run.font.size  = Pt(11)
    run.font.bold  = bold
    run.font.color.rgb = color or LIGHT
    return p


def bullet(doc, text, sub=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.5 if sub else 0.3)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    run = p.add_run(text)
    run.font.name  = "Segoe UI"
    run.font.size  = Pt(11)
    run.font.color.rgb = LIGHT
    return p


def tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("Tip:  ")
    run.font.name  = "Segoe UI"
    run.font.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = GOLD
    run2 = p.add_run(text)
    run2.font.name  = "Segoe UI"
    run2.font.size  = Pt(11)
    run2.font.color.rgb = LIGHT
    return p


def shortcut_table(doc, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Button / Action"
    hdr[1].text = "What it does"
    for cell in hdr:
        set_cell_bg(cell, "0C0C14")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = GOLD
            run.font.bold = True
            run.font.name = "Segoe UI"
            run.font.size = Pt(10)

    for i, (action, desc) in enumerate(rows):
        r = table.rows[i + 1].cells
        r[0].text = action
        r[1].text = desc
        bg = "111120" if i % 2 == 0 else "0C0C14"
        for cell in r:
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = LIGHT
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(10)

    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.0)
    doc.add_paragraph()


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading("ProfTracker", 0)
    for run in title.runs:
        run.font.color.rgb = GOLD
        run.font.name = "Segoe UI"
        run.font.size = Pt(28)
        run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Marvel Rivals Hero Proficiency Tracker — User Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = GRAY
        run.font.size = Pt(12)
        run.font.name = "Segoe UI"
    doc.add_paragraph()

    # ── 1. Overview ────────────────────────────────────────────────────────
    heading(doc, "1.  Overview")
    body(doc,
         "ProfTracker lets you record and track the proficiency level of every hero "
         "in Marvel Rivals — whether by scanning the game directly, entering data "
         "manually, or importing a filled-in Excel spreadsheet. All data is stored "
         "locally; no account or internet connection is required to use the app.")

    # ── 2. Installation ────────────────────────────────────────────────────
    heading(doc, "2.  Installation")
    body(doc,
         "ProfTracker is distributed as a pre-built folder — no Python or coding "
         "knowledge required. Download the latest ProfTracker-windows.zip from the "
         "GitHub Releases page, unzip it anywhere, and run ProfTracker.exe.")
    body(doc, "Depending on which features you want to use, two optional system "
              "installs may be needed:", color=LIGHT)

    heading(doc, "2a.  Tesseract OCR  (required for screen scanning)", level=2)
    body(doc,
         "The screen scanner reads hero data directly from the game window using "
         "OCR (text recognition). Tesseract is the engine that powers this.")
    bullet(doc, "Download the installer from: github.com/UB-Mannheim/tesseract/wiki")
    bullet(doc, "Run the installer and accept the default install path "
                "(C:\\Program Files\\Tesseract-OCR\\).")
    bullet(doc, "No reboot required.")
    body(doc, "Without Tesseract, Manual Entry and Excel Import still work fine — "
              "only the screen scanner is unavailable.", color=GRAY)

    heading(doc, "2b.  Interception Driver  (required for Auto Scan only)", level=2)
    body(doc,
         "Auto Scan drives the game automatically using synthetic mouse and keyboard "
         "input. This requires the Interception kernel driver.")
    bullet(doc, "Download Interception from: github.com/oblitum/Interception")
    bullet(doc, "Run the installer as Administrator.")
    bullet(doc, "Reboot your PC after installation.")
    body(doc, "The Interception driver is only needed for Auto Scan. Manual Session, "
              "Manual Entry, and Excel Import do not require it.", color=GRAY)
    tip(doc, "If you only plan to use Manual Entry or Import Excel, you can skip "
             "both installs entirely and just run ProfTracker.exe.")

    # ── 3. Navigation ──────────────────────────────────────────────────────
    heading(doc, "3.  Navigation")
    body(doc,
         "The top bar contains three tabs. Click any tab to switch pages. "
         "The current version number is shown on the right side of the bar.")
    shortcut_table(doc, [
        ("SCAN",        "Set up a scan session or run an automatic scan of the game."),
        ("HEROES",      "Browse, sort, filter, and inspect all recorded heroes."),
        ("XP PROGRESS", "View a chart of XP earned over time for selected heroes."),
    ])

    # ── 4. Scanning Heroes ─────────────────────────────────────────────────
    heading(doc, "4.  Scanning Heroes")
    body(doc, "Scanning reads hero data directly from the Marvel Rivals game window "
              "using on-screen text recognition (OCR). There are two scan modes.")
    body(doc, "⚠  Required game settings before scanning:", bold=True, color=GOLD)
    bullet(doc, "Display Mode: Windowed (not Fullscreen or Borderless Windowed).")
    bullet(doc, "Resolution: 1920 × 1080 (1080p).")
    body(doc, "Scanning at any other resolution or display mode will produce incorrect "
              "results or fail entirely. Set these in Marvel Rivals → Settings → Video "
              "before starting a session.", color=GRAY)

    heading(doc, "4a.  Selecting the Game Window", level=2)
    body(doc, "Before scanning, click Select Window and then click anywhere on the "
              "Marvel Rivals window. The status bar will confirm the window was found.")

    heading(doc, "4b.  Manual Session", level=2)
    body(doc, "In a manual session you navigate to each hero's Proficiency tab yourself "
              "and press 2 (or click 'Capture Proficiency') to record that hero.")
    bullet(doc, "Click Start Session — the window shrinks to a compact overlay.")
    bullet(doc, "In-game, open a hero's profile and go to their Proficiency tab.")
    bullet(doc, "Select the hero name in the dropdown, then press 2 or click the button.")
    bullet(doc, "Repeat for each hero. Click Finish when done.")
    tip(doc, "The app stays on top of the game window during a manual session.")

    heading(doc, "4c.  Auto Scan", level=2)
    body(doc, "Auto Scan drives the game automatically, cycling through every hero "
              "and capturing each one.")
    bullet(doc, "Click Auto Scan to begin.")
    bullet(doc, "Move your mouse once and press any key to register input devices.")
    bullet(doc, "You have 4 seconds to switch to the game window before scanning starts.")
    bullet(doc, "Click Cancel Scan at any time to stop early.")
    tip(doc, "Auto Scan works best on the Heroes → Proficiency view with the game "
             "at its default layout. Avoid using the mouse or keyboard during the scan.")

    # ── 5. Manual Entry ────────────────────────────────────────────────────
    heading(doc, "5.  Manual Entry")
    body(doc, "You can add or update a hero's data without opening the game at all.")
    bullet(doc, "In the HEROES tab, click + Manual Entry.")
    bullet(doc, "Select a hero from the dropdown. If the hero already has data saved, "
                "the form pre-fills with their current values.")
    bullet(doc, "Set the Level and Current XP. Tick Max Level if the hero is LV 70.")
    bullet(doc, "Click Save. The hero grid updates immediately.")
    tip(doc, "XP Required is calculated automatically from the level — you don't need to enter it.")

    # ── 6. Import from Excel ───────────────────────────────────────────────
    heading(doc, "6.  Importing from Excel")
    body(doc, "If you want to enter data for many heroes at once, use the Excel import.")

    heading(doc, "6a.  Getting the Template", level=2)
    bullet(doc, "Click Get Template in the HEROES toolbar.")
    bullet(doc, "Save the file anywhere on your computer.")
    bullet(doc, "Open it in Microsoft Excel or Google Sheets.")

    heading(doc, "6b.  Filling in the Template", level=2)
    body(doc, "The 'Heroes' sheet contains every hero pre-filled at Level 1.")
    bullet(doc, "Change the Level column to the hero's current level.")
    bullet(doc, "Enter the Current XP value shown on the Proficiency tab in-game.")
    bullet(doc, "Set Max Level to Yes for any hero at LV 70.")
    bullet(doc, "Leave the Hero Name and Role columns as-is.")
    tip(doc, "Rows you don't change (left at Level 1, XP 0) will still be imported "
             "and will overwrite any existing data for that hero. Only edit rows for "
             "heroes whose data you actually know.")

    heading(doc, "6c.  Importing the File", level=2)
    bullet(doc, "Click Import Excel in the HEROES toolbar.")
    bullet(doc, "Select your filled-in spreadsheet.")
    bullet(doc, "A summary will appear showing how many heroes were imported and "
                "any rows that were skipped (with the reason).")

    # ── 7. Hero Browser ────────────────────────────────────────────────────
    heading(doc, "7.  Hero Browser")
    body(doc, "The HEROES tab shows all recorded heroes as cards. "
              "Click any card to see full detail.")
    shortcut_table(doc, [
        ("Role filter",          "Show only Vanguards, Duelists, Strategists, or all heroes."),
        ("Sort: Closest to Lord",
         "Heroes nearest LV 20 (Lord rank) appear first — useful for grinding."),
        ("Sort: Closest to Champion",
         "Heroes nearest LV 50 (Champion rank) appear first."),
        ("Sort: Alphabetical",   "A–Z by hero name."),
        ("Sort: Level",          "Lowest to highest level (flip with Asc/Desc toggle)."),
        ("↑ Asc / ↓ Desc",       "Toggle sort direction."),
        ("Export CSV",           "Save all currently visible heroes to a CSV file."),
    ])

    # ── 8. Hero Detail ─────────────────────────────────────────────────────
    heading(doc, "8.  Hero Detail")
    body(doc, "Click any hero card to open the detail panel.")
    bullet(doc, "Shows the hero icon (animated for Champion-rank heroes).")
    bullet(doc, "Current level, XP progress bar, and total XP earned.")
    bullet(doc, "A second progress bar shows overall progress toward Champion (LV 50).")
    bullet(doc, "MAX LEVEL is displayed for LV 70 heroes instead of the XP bar.")

    # ── 9. XP Progress ─────────────────────────────────────────────────────
    heading(doc, "9.  XP Progress Chart")
    body(doc, "The XP PROGRESS tab plots total XP earned over time for your heroes.")
    bullet(doc, "Use the time range selector (7 / 14 / 30 / 90 days, or All Time).")
    bullet(doc, "Check or uncheck individual heroes to show or hide their line.")
    tip(doc, "Data points are added each time you scan or manually update a hero, "
             "so scan regularly to get meaningful progress trends.")

    # ── 10. Updates ────────────────────────────────────────────────────────
    heading(doc, "10.  Updates")
    body(doc, "ProfTracker checks for updates automatically each time it launches "
              "(requires an internet connection).")
    bullet(doc, "If a newer version is available, a gold banner appears at the top of the window.")
    bullet(doc, "Click the Download button in the banner to open the releases page in your browser.")
    bullet(doc, "Download the new zip, unzip it, and replace your old ProfTracker folder.")

    doc.save("How to Use ProfTracker.docx")
    print("Saved: How to Use ProfTracker.docx")


if __name__ == "__main__":
    build()
